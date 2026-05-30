from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("java_intern_interview_answers_knowflow.docx")


QA_ITEMS = [
    (
        "1. 项目中用户登录后使用 JWT 进行认证。请你说说 JWT 的组成、服务端如何校验 Token，以及相比 Session 登录有什么优缺点？",
        "面试回答：JWT 一般由 Header、Payload、Signature 三部分组成。Header 描述签名算法，Payload 放用户标识、角色、过期时间等声明，Signature 是服务端用密钥对前两部分签名后的结果。服务端收到请求后，通常从 Authorization 头取出 Bearer Token，校验签名是否正确、是否过期、是否被吊销，再把用户身份放到 Spring Security 的上下文里。相比 Session，JWT 的优势是无状态，适合前后端分离、网关转发和水平扩容，服务端不用为每个用户保存会话。但它也有缺点：一旦签发，在过期前很难天然失效，所以项目里最好结合 Redis 做黑名单或刷新 Token；同时 Payload 只是 Base64 编码，不应该放密码、密钥这类敏感信息。",
    ),
    (
        "2. 如果用户上传一个 500MB 的 PDF，后端为什么不能直接一次性接收完整文件？这个项目里的分片上传、断点续传大概可以怎么设计？",
        "面试回答：500MB 文件如果一次性上传，会占用较长连接时间和较多内存、磁盘缓冲，也容易因为网络抖动导致整次上传失败。更合理的方式是分片上传。前端先计算文件 MD5，把文件按固定大小切片，例如每片 5MB，上传时带上 fileMd5、chunkIndex、totalChunks 等信息。后端每收到一片，就把它临时写入 MinIO 或本地临时区，同时用 Redis Bitmap 或 Set 记录哪些分片已完成。断点续传时，前端先查询上传状态，只补传缺失分片。所有分片完成后，调用 merge 接口合并文件，生成文件记录，再发送异步任务做解析和向量化。这样可以降低失败成本，也能提升大文件上传体验。",
    ),
    (
        "3. 项目中上传文件后需要解析、切块、向量化、写入 Elasticsearch。你会把这些步骤放在一次 HTTP 请求里同步完成吗？为什么？",
        "面试回答：我不会全部放在一次 HTTP 请求里同步完成。上传文件是一个用户交互操作，应该尽快返回“上传成功或已进入处理队列”。解析、切块、Embedding 向量化、写 Elasticsearch 都可能耗时很长，还依赖外部模型服务和 ES，如果同步执行，接口容易超时，用户体验也差。更好的设计是上传合并完成后，先在数据库中创建文件和处理任务记录，状态为 PENDING，然后发送 Kafka 消息，由消费者异步处理。处理过程中不断更新任务状态，例如 PARSING、VECTORIZING、INDEXING、SUCCESS、FAILED。前端可以轮询或通过 WebSocket 查看进度。这样能解耦上传链路和计算链路，也方便失败重试、限流和横向扩容。",
    ),
    (
        "4. Kafka 在这个项目中适合承担什么角色？如果文件上传成功后发送消息失败，会造成什么问题？你会怎么保证可靠性？",
        "面试回答：Kafka 在这个项目里适合承担异步任务队列的角色，例如文件上传完成后触发解析、切块、向量化、索引写入等后台任务。它可以把用户请求和重计算任务解耦，削峰填谷。如果文件上传成功但发送 Kafka 消息失败，数据库里可能已有文件记录，MinIO 里也有文件，但后续解析任务没有被执行，最终表现为文档一直不可检索。可靠性上，我会先把文件和任务状态落库，再发送消息；如果发送失败，任务仍然保留为 PENDING，可以由定时补偿任务扫描并重新投递。消费者处理时要做幂等，例如按 fileMd5 或 taskId 判断是否已处理，避免重复消息导致重复索引。必要时还可以使用事务消息或 outbox 模式进一步增强一致性。",
    ),
    (
        "5. 项目使用 Redis 记录上传分片状态。为什么 Redis 适合这个场景？如果 Redis 数据过期或丢失，会对断点续传造成什么影响？",
        "面试回答：Redis 适合记录分片状态，主要因为它读写快、结构灵活、天然支持过期时间。上传分片状态是典型的临时数据，不一定需要长期保存在 MySQL 中。比如可以用 Redis Bitmap 记录某个 fileMd5 下每个分片是否上传完成，查询缺失分片非常快，也节省空间。设置 TTL 可以自动清理长时间未完成的上传任务，避免垃圾数据堆积。如果 Redis 数据过期或丢失，断点续传就无法准确知道哪些分片已上传，前端可能需要重新上传全部分片，或者后端要从临时存储中重新扫描已存在分片来恢复状态。因此 Redis 不能作为唯一可信事实，关键文件元数据和最终状态仍应落库，Redis 更适合做加速和临时状态管理。",
    ),
    (
        "6. Spring Boot 中 @Transactional 的作用是什么？在删除文档时，如果数据库记录删除成功，但 MinIO 文件删除失败，会出现什么一致性问题？",
        "面试回答：@Transactional 用来声明事务边界，保证同一个事务里的数据库操作要么一起提交，要么一起回滚。比如删除文档时，可能要删文件记录、chunk 记录、向量记录等数据库数据，这些操作可以放在一个事务里。但 MinIO 和 Elasticsearch 不属于同一个数据库事务，@Transactional 无法自动回滚这些外部系统操作。如果数据库记录删成功，但 MinIO 文件删除失败，就会产生孤儿文件，占用存储空间；如果 ES 索引没删干净，还可能出现搜索到已删除文档的情况。实践中可以采用最终一致性：数据库先标记 DELETING，执行外部删除，成功后改为 DELETED；失败则记录补偿任务，后台重试。对用户侧可以先隐藏该文档，避免影响业务体验。",
    ),
    (
        "7. JPA 的 JpaRepository 给开发带来了什么便利？如果知识库文档列表需要按用户、组织标签、公开状态过滤，你会如何设计查询？",
        "面试回答：JpaRepository 提供了常用 CRUD、分页、排序和按方法名派生查询的能力，能减少大量样板代码。对于知识库文档列表，我会先明确权限规则：用户可以看到自己上传的文档、公开文档，以及自己组织标签范围内的非公开文档。简单场景可以在 Repository 中写 JPQL，例如按 userId、isPublic、orgTag in orgTags 组合过滤，并配合 Pageable 做分页。复杂场景下，我更倾向使用 Specification、Criteria API 或 QueryDSL 动态构造条件，避免方法名过长。还要注意索引设计，例如 userId、isPublic、orgTag、createdTime 建索引。返回给前端时不要直接暴露实体，最好转成 DTO，避免懒加载、字段泄露和序列化问题。",
    ),
    (
        "8. 项目中有用户、管理员、组织标签等权限模型。Spring Security 通常如何实现接口级权限控制？JWT 里的角色信息能完全信任吗？",
        "面试回答：Spring Security 通常通过过滤器链完成认证，再通过注解或配置完成授权。比如 JWT 过滤器解析 Token，校验通过后构造 Authentication，放入 SecurityContext；接口上可以使用 @PreAuthorize(\"hasRole('ADMIN')\") 做角色控制，也可以在业务层基于 userId、orgTag 做资源级权限校验。JWT 里的角色信息不能脱离签名校验直接信任，因为 Payload 本身可读可改，只有签名正确才说明没被篡改。即使签名有效，也要考虑角色变更后的及时性：如果用户被降权，但旧 Token 还没过期，可能继续访问管理员接口。因此项目中可以缩短 Access Token 有效期，使用 Refresh Token，或在 Redis 中维护用户权限版本、Token 黑名单，关键操作再查库确认。",
    ),
    (
        "9. 项目中支持 WebSocket 流式聊天。WebSocket 和普通 HTTP 请求有什么区别？为什么 AI 对话场景更适合使用 WebSocket 或 SSE？",
        "面试回答：普通 HTTP 是典型的请求响应模型，客户端发一次请求，服务端返回一次结果。WebSocket 则是在一次握手后建立长连接，客户端和服务端可以双向实时通信。AI 对话场景中，大模型通常是流式生成内容，如果等全部生成完再一次性返回，用户会感觉等待时间很长。使用 WebSocket 或 SSE 可以边生成边推送，用户能实时看到回答过程。WebSocket 更适合双向互动，例如用户中途停止生成、服务端推送状态、断线重连恢复；SSE 更轻量，适合服务端单向流式输出。这个项目有聊天会话、生成状态、引用预览等能力，用 WebSocket 可以更自然地管理连接、消息和实时状态。",
    ),
    (
        "10. 如果多个用户同时对话，后端如何区分每个用户的 WebSocket 连接？你会如何设计连接管理和断线重连？",
        "面试回答：多个用户同时对话时，后端需要在 WebSocket 建连阶段完成身份认证。常见做法是客户端先通过 HTTP 获取一个短期 websocket-token，连接时带上 token，服务端校验后解析出 userId、conversationId 等信息。连接管理可以用 ConcurrentHashMap 保存 userId 或 sessionId 到 WebSocketSession 的映射，如果一个用户有多端登录，也可以保存为列表。断线时要清理连接，避免内存泄漏。断线重连方面，可以为每次 AI 生成分配 generationId，生成状态和已输出片段保存在 Redis 或数据库中；客户端重连后查询 active-generation，再继续接收或拉取缺失内容。这样可以避免网络波动导致回答丢失，也方便实现停止生成和重复连接控制。",
    ),
    (
        "11. 文档解析后要切分成多个文本块。假设一个大文档被切成 1 万个 chunk，批量写入数据库或 Elasticsearch 时要注意哪些性能问题？",
        "面试回答：1 万个 chunk 如果逐条写数据库或 ES，会产生大量网络往返和事务开销，性能会很差。数据库侧应该使用批量插入，控制 batch size，例如每 500 或 1000 条提交一次，避免单个事务过大导致锁时间长、内存占用高。Elasticsearch 侧应使用 Bulk API 批量写入，同样要控制每批文档数量和请求体大小，处理部分失败结果，失败项单独重试。向量字段通常很大，还要关注 ES mapping、refresh 策略和索引写入压力，必要时降低刷新频率。应用层也要避免一次性把所有 chunk、embedding 全部放内存，可以流式处理或分批处理。同时要设计幂等键，例如 fileMd5 + chunkIndex，防止重试时重复写入。",
    ),
    (
        "12. 项目中有每日额度、Token 使用量、限流配置。如果多个请求并发扣减同一个用户额度，可能出现什么并发问题？如何解决？",
        "面试回答：并发扣减额度时，最典型的问题是超扣或少扣。比如用户剩余 1000 token，同时来了两个请求，每个都读取到 1000，然后分别扣 800，最后可能都成功，实际超出了额度。解决方式要保证扣减操作的原子性。数据库方案可以使用乐观锁，在额度记录上加 version 字段，更新时带上 version，不匹配就重试；也可以用条件更新，例如 update quota set remain = remain - ? where user_id = ? and remain >= ?。Redis 方案可以用 Lua 脚本把判断和扣减放在一个原子操作里，适合高频限流。对于最终账单，最好把使用流水落库，余额可以异步对账修正。关键是不能只在 Java 内存里判断，否则多实例部署时会失效。",
    ),
    (
        "13. Java 中 HashMap、ConcurrentHashMap 有什么区别？如果用内存 Map 保存在线聊天连接，为什么普通 HashMap 可能有风险？",
        "面试回答：HashMap 不是线程安全的，多个线程同时读写时可能出现数据覆盖、读到脏数据，甚至在旧版本中扩容时出现链表异常问题。ConcurrentHashMap 是并发安全的 Map，它通过更细粒度的同步和 CAS 等机制保证并发访问下的数据一致性，同时性能比整表加锁更好。在 WebSocket 场景中，连接建立、消息发送、连接关闭通常由不同线程触发，如果用普通 HashMap 保存 session 映射，就可能出现刚加入的连接被覆盖、关闭时删除错误、遍历发送消息时并发修改异常等问题。因此在线连接表应该使用 ConcurrentHashMap，并且在多实例部署时还要意识到内存 Map 只保存本机连接。如果要跨节点路由消息，需要引入 Redis、消息队列或网关层做协调。",
    ),
    (
        "14. 项目中需要调用 DeepSeek 或其他大模型 API。如果外部 API 响应很慢或失败，后端应该如何处理超时、重试和降级？",
        "面试回答：调用外部大模型 API 一定要设置超时，包括连接超时、读取超时和整体请求超时，不能无限等待。失败处理要区分类型：网络抖动、HTTP 5xx 可以有限重试；参数错误、鉴权失败、余额不足这类 4xx 一般不应该盲目重试。重试要有次数上限和退避策略，避免把外部服务打得更慢。降级方面，可以给用户返回明确提示，比如“模型服务暂时不可用，请稍后重试”；如果项目配置了多个模型供应商，也可以通过路由切换到备用供应商。对于流式聊天，还要在中途异常时推送错误事件并记录 generation 状态。除此之外，还应做熔断、限流、日志追踪和调用耗时监控，方便定位是网络、模型服务还是自身线程池压力问题。",
    ),
    (
        "15. 请你设计一个“删除知识库文档”的完整后端流程：需要删除哪些数据？涉及数据库、MinIO、Elasticsearch、缓存时，如何尽量保证最终一致性？",
        "面试回答：删除知识库文档时，我会先做权限校验，确认当前用户是文档所有者或管理员，并且有对应组织权限。然后在数据库事务中把文件记录状态改为 DELETING，或删除文件元数据、chunk 信息、向量任务记录等业务数据。接着删除 Elasticsearch 中该 fileMd5 关联的索引文档，删除 MinIO 中的原始文件和预览文件，再清理 Redis 中的上传状态、检索缓存或文档权限缓存。由于数据库、ES、MinIO、Redis 不能放进一个本地事务，我会采用最终一致性：核心数据库状态先变更，外部删除失败时记录补偿任务，后台定时重试；接口对用户立即隐藏 DELETING 文档，避免继续被检索或下载。所有删除操作都要幂等，重复执行不能报不可恢复错误，比如文件不存在应视为删除成功。",
    ),
]


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_spacing(paragraph, before=0, after=6, line=300):
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"), str(after * 20))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 18, 10),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 14, 7),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    title_style = styles["Title"]
    title_style.font.name = "Microsoft YaHei"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(6)
    title_style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "KnowFlow 项目 Java 实习生面试题参考回答"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    set_east_asia_font(header.runs[0])

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        set_east_asia_font(run)


def add_title(doc):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KnowFlow 项目 Java 实习生面试题与参考回答")
    set_east_asia_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("面向互联网公司 Java 技术面试，围绕 RAG 知识库项目的真实工程场景展开")
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    set_paragraph_spacing(subtitle, before=0, after=14, line=300)

    note = doc.add_paragraph()
    run = note.add_run(
        "使用建议：回答时不要死背概念，优先按“场景背景 -> 技术原理 -> 项目落地 -> 风险与优化”的顺序组织语言。"
    )
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    run.font.bold = True
    set_paragraph_spacing(note, before=0, after=12, line=300)


def add_questions(doc):
    for question, answer in QA_ITEMS:
        q = doc.add_paragraph(style="Heading 1")
        q_run = q.add_run(question)
        set_east_asia_font(q_run)
        q_run.font.bold = True

        a = doc.add_paragraph()
        a.paragraph_format.first_line_indent = Inches(0.22)
        a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        a_run = a.add_run(answer)
        set_east_asia_font(a_run)
        a_run.font.size = Pt(11)
        set_paragraph_spacing(a, before=0, after=8, line=300)


def main():
    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_questions(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
